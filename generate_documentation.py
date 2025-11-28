"""
Script to generate comprehensive DOCX documentation for PictoBlox Python Backend
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

def add_heading_with_style(doc, text, level=1):
    """Add heading with custom styling"""
    heading = doc.add_heading(text, level=level)
    heading.style.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    return heading

def add_code_block(doc, code_text):
    """Add code block with monospace font"""
    para = doc.add_paragraph()
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    para.style = 'No Spacing'

def create_documentation():
    """Create comprehensive documentation"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_heading('PictoBlox Python Backend', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Project Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)
    subtitle_format.italic = True
    
    doc.add_paragraph()
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        ('1. Project Overview', 1),
        ('2. Architecture & Technology Stack', 2),
        ('3. Project Structure', 2),
        ('4. Database Models', 2),
        ('5. API Endpoints', 2),
        ('6. Authentication & Security', 2),
        ('7. Code Execution & Sandbox', 2),
        ('8. Installation & Setup', 2),
        ('9. Configuration', 2),
        ('10. Usage Examples', 2),
        ('11. Development Guidelines', 2),
        ('12. Deployment', 2),
    ]
    
    for item, level in toc_items:
        para = doc.add_paragraph(item, style='List Number' if level == 1 else 'List Bullet 2')
    
    doc.add_page_break()
    
    # 1. Project Overview
    doc.add_heading('1. Project Overview', 1)
    doc.add_paragraph(
        'PictoBlox Python Backend is a comprehensive FastAPI-based REST API backend for a Blockly-based '
        'visual programming platform. Similar to Scratch or PictoBlox, this platform allows users to '
        'create interactive projects using visual block-based programming, which can be converted to '
        'Python or JavaScript code and executed in a sandboxed environment.'
    )
    
    doc.add_heading('1.1 Key Features', 2)
    features = [
        'User authentication and authorization with JWT tokens',
        'Project management (create, read, update, delete, duplicate)',
        'Sprite and costume management for visual programming',
        'Backdrop and stage settings management',
        'Variable and list management (global and sprite-specific)',
        'Code execution in sandboxed environment (Python and JavaScript)',
        'Asset management (sprites, sounds, images, videos)',
        'Project sharing and collaboration',
        'Library of pre-made sprites and backdrops',
        'Execution logging and history',
        'Public project discovery',
        'Extension system for custom blocks'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('1.2 Technology Stack', 2)
    tech_stack = [
        ('Backend Framework', 'FastAPI 0.104.1'),
        ('Database ORM', 'SQLAlchemy 2.0.23'),
        ('Database', 'PostgreSQL (Supabase) / SQLite / MySQL'),
        ('Authentication', 'JWT (python-jose), bcrypt'),
        ('Validation', 'Pydantic 2.9.2'),
        ('Server', 'Uvicorn'),
        ('Image Processing', 'Pillow (PIL)'),
        ('Code Execution', 'Python subprocess, Docker (optional)')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    
    for component, tech in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = tech
    
    # 2. Architecture & Technology Stack
    doc.add_page_break()
    doc.add_heading('2. Architecture & Technology Stack', 1)
    
    doc.add_heading('2.1 System Architecture', 2)
    doc.add_paragraph(
        'The application follows a layered architecture pattern:'
    )
    layers = [
        ('API Layer', 'FastAPI routes and endpoints handling HTTP requests'),
        ('Business Logic Layer', 'CRUD operations and business rules in crud.py'),
        ('Data Access Layer', 'SQLAlchemy ORM models and database interactions'),
        ('Authentication Layer', 'JWT-based authentication and authorization'),
        ('Execution Layer', 'Sandboxed code execution environment')
    ]
    
    for layer, description in layers:
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(f'{layer}: ')
        run.bold = True
        para.add_run(description)
    
    doc.add_heading('2.2 Database Architecture', 2)
    doc.add_paragraph(
        'The database uses PostgreSQL (via Supabase) with the following main entities:'
    )
    entities = [
        'Users - User accounts and profiles',
        'Projects - Visual programming projects with Blockly workspace data',
        'Sprites - Visual characters/objects in projects',
        'Costumes - Different appearances for sprites',
        'Backdrops - Stage backgrounds',
        'StageSettings - Stage configuration',
        'SpriteVariables - Variables (global or sprite-specific)',
        'SpriteLists - Lists (global or sprite-specific)',
        'Assets - General asset storage (sprites, sounds, etc.)',
        'ExecutionLogs - Code execution history',
        'ProjectShares - Project sharing and collaboration',
        'Extensions - Custom Blockly extensions',
        'LibrarySprites - Pre-made sprite library',
        'LibraryBackdrops - Pre-made backdrop library',
        'HardwareDevices - Connected hardware devices'
    ]
    
    for entity in entities:
        doc.add_paragraph(entity, style='List Bullet')
    
    # 3. Project Structure
    doc.add_page_break()
    doc.add_heading('3. Project Structure', 1)
    
    structure = """
pictoblox-python-backend/
├── main.py              # FastAPI application and API endpoints
├── models.py            # SQLAlchemy database models
├── schemas.py           # Pydantic schemas for request/response validation
├── crud.py              # CRUD operations for database models
├── database.py          # Database configuration and session management
├── auth.py              # Authentication and security functions
├── sandbox.py           # Code execution sandbox
├── requirements.txt     # Python dependencies
└── blockly_platform.db  # SQLite database (if used)
"""
    
    add_code_block(doc, structure)
    
    doc.add_heading('3.1 File Descriptions', 2)
    
    file_descriptions = [
        ('main.py', 'Main FastAPI application file containing all API endpoints. Handles authentication, project management, sprite management, code execution, and more.'),
        ('models.py', 'SQLAlchemy ORM models defining database schema. Includes User, Project, Sprite, Costume, Backdrop, and other models.'),
        ('schemas.py', 'Pydantic models for request/response validation. Ensures data integrity and provides automatic API documentation.'),
        ('crud.py', 'Database CRUD (Create, Read, Update, Delete) operations. Contains all database interaction logic separated from API endpoints.'),
        ('database.py', 'Database configuration, connection management, and session factory. Supports PostgreSQL, MySQL, and SQLite.'),
        ('auth.py', 'Authentication and authorization logic. JWT token generation/validation, password hashing, and user verification.'),
        ('sandbox.py', 'Code execution sandbox for running user-generated Python and JavaScript code safely with timeout and resource limits.')
    ]
    
    for filename, description in file_descriptions:
        para = doc.add_paragraph()
        run = para.add_run(f'{filename}: ')
        run.bold = True
        para.add_run(description)
    
    # 4. Database Models
    doc.add_page_break()
    doc.add_heading('4. Database Models', 1)
    
    doc.add_heading('4.1 Core Models', 2)
    
    models_info = [
        ('User', 'User accounts with authentication. Fields: id, username, email, hashed_password, role, is_active, is_verified, avatar_url, bio, created_at, updated_at'),
        ('Project', 'Visual programming projects. Stores Blockly workspace XML/JSON, generated code, metadata, statistics (views, likes, forks), and version control.'),
        ('Sprite', 'Visual characters/objects. Properties: position (x, y), direction, size, visibility, rotation_style, layer_order, costumes, variables, lists.'),
        ('Costume', 'Sprite appearances. Stores image data (URL or base64), dimensions, rotation center, and order.'),
        ('Backdrop', 'Stage backgrounds. Similar to costumes but for the stage.'),
        ('StageSetting', 'Stage configuration: dimensions, current backdrop, audio settings, video settings, FPS, coordinate system.'),
        ('SpriteVariable', 'Variables for sprites or global project variables. Can be cloud variables, visible on stage, with position.'),
        ('SpriteList', 'Lists for sprites or global project lists. Stores array of items, can be displayed on stage.'),
        ('Asset', 'General asset storage for sprites, sounds, images, videos. Supports file URLs or base64 data.'),
        ('ExecutionLog', 'Code execution history. Stores code, output, errors, execution time, and metadata.'),
        ('ProjectShare', 'Project sharing and collaboration. Defines permissions (can_edit, can_view, can_execute).'),
        ('Extension', 'Custom Blockly extensions. Contains block definitions and generator code.'),
        ('LibrarySprite', 'Pre-made sprites available in library. Includes categories, tags, download counts.'),
        ('LibraryBackdrop', 'Pre-made backdrops available in library.'),
        ('HardwareDevice', 'Connected hardware devices (Arduino, ESP32, etc.) with connection details and capabilities.')
    ]
    
    for model_name, description in models_info:
        para = doc.add_paragraph()
        run = para.add_run(f'{model_name}: ')
        run.bold = True
        para.add_run(description)
    
    doc.add_heading('4.2 Relationships', 2)
    relationships = [
        'User has many Projects (one-to-many)',
        'User has many ExecutionLogs (one-to-many)',
        'Project belongs to User (many-to-one)',
        'Project has many Sprites (one-to-many)',
        'Project has many Backdrops (one-to-many)',
        'Project has one StageSetting (one-to-one)',
        'Project has many Assets (one-to-many)',
        'Project has many ExecutionLogs (one-to-many)',
        'Sprite has many Costumes (one-to-many)',
        'Sprite has many Variables (one-to-many)',
        'Sprite has many Lists (one-to-many)',
        'StageSetting references one Backdrop (many-to-one)',
        'ProjectShare links Project to User (many-to-many)'
    ]
    
    for rel in relationships:
        doc.add_paragraph(rel, style='List Bullet')
    
    # 5. API Endpoints
    doc.add_page_break()
    doc.add_heading('5. API Endpoints', 1)
    
    doc.add_heading('5.1 Authentication Endpoints', 2)
    auth_endpoints = [
        ('POST /api/v1/register', 'Register a new user. Requires: username, email, password. Returns: User object.'),
        ('POST /api/v1/token', 'Login and get access token. Uses OAuth2PasswordRequestForm. Returns: access_token, token_type.'),
        ('GET /api/v1/users/me', 'Get current user information. Requires: Bearer token. Returns: User object.'),
        ('PUT /api/v1/users/me', 'Update current user profile. Requires: Bearer token. Returns: Updated User object.')
    ]
    
    for endpoint, description in auth_endpoints:
        para = doc.add_paragraph()
        run = para.add_run(f'{endpoint}: ')
        run.bold = True
        para.add_run(description)
    
    doc.add_heading('5.2 Project Management Endpoints', 2)
    project_endpoints = [
        ('POST /api/v1/projects', 'Create a new project. Requires: Bearer token. Returns: Project object.'),
        ('GET /api/v1/projects', 'List all projects for current user. Supports pagination (skip, limit).'),
        ('GET /api/v1/projects/{project_id}', 'Get a specific project. Checks authorization.'),
        ('PUT /api/v1/projects/{project_id}', 'Update a project. Only owner can update.'),
        ('DELETE /api/v1/projects/{project_id}', 'Delete a project. Only owner can delete.'),
        ('POST /api/v1/projects/{project_id}/duplicate', 'Duplicate/clone a project.'),
        ('GET /api/v1/projects/list/public', 'List all public projects. No authentication required.')
    ]
    
    for endpoint, description in project_endpoints:
        para = doc.add_paragraph()
        run = para.add_run(f'{endpoint}: ')
        run.bold = True
        para.add_run(description)
    
    doc.add_heading('5.3 Sprite Management Endpoints', 2)
    sprite_endpoints = [
        ('POST /api/v1/sprites', 'Create a new sprite in a project.'),
        ('GET /api/v1/projects/{project_id}/sprites', 'Get all sprites for a project. Can include costumes.'),
        ('GET /api/v1/sprites/{sprite_id}', 'Get a single sprite with optional related data.'),
        ('PUT /api/v1/sprites/{sprite_id}', 'Update a sprite.'),
        ('DELETE /api/v1/sprites/{sprite_id}', 'Delete a sprite.'),
        ('POST /api/v1/sprites/{sprite_id}/duplicate', 'Duplicate/clone a sprite with all costumes.'),
        ('PUT /api/v1/projects/{project_id}/sprites/reorder', 'Reorder sprite layers (z-index).'),
        ('PUT /api/v1/sprites/{sprite_id}/front', 'Bring sprite to front (top layer).'),
        ('PUT /api/v1/sprites/{sprite_id}/back', 'Send sprite to back (bottom layer).'),
        ('PUT /api/v1/sprites/{sprite_id}/move', 'Move sprite to position.'),
        ('PUT /api/v1/sprites/{sprite_id}/rotate', 'Rotate sprite to direction.'),
        ('PUT /api/v1/sprites/{sprite_id}/size', 'Change sprite size.'),
        ('PUT /api/v1/sprites/{sprite_id}/visibility', 'Show/hide sprite.')
    ]
    
    for endpoint, description in sprite_endpoints:
        para = doc.add_paragraph()
        run = para.add_run(f'{endpoint}: ')
        run.bold = True
        para.add_run(description)
    
    doc.add_heading('5.4 Costume Management Endpoints', 2)
    costume_endpoints = [
        ('POST /api/v1/costumes', 'Create a new costume.'),
        ('POST /api/v1/costumes/upload', 'Upload costume image file. Accepts multipart/form-data.'),
        ('GET /api/v1/sprites/{sprite_id}/costumes', 'Get all costumes for a sprite.'),
        ('PUT /api/v1/costumes/{costume_id}', 'Update a costume.'),
        ('DELETE /api/v1/costumes/{costume_id}', 'Delete a costume.'),
        ('PUT /api/v1/sprites/{sprite_id}/costume', 'Set active costume for sprite.'),
        ('POST /api/v1/costumes/{costume_id}/duplicate', 'Duplicate a costume.')
    ]
    
    for endpoint, description in costume_endpoints:
        para = doc.add_paragraph()
        run = para.add_run(f'{endpoint}: ')
        run.bold = True
        para.add_run(description)
    
    doc.add_heading('5.5 Backdrop Management Endpoints', 2)
    backdrop_endpoints = [
        ('POST /api/v1/backdrops', 'Create a new backdrop.'),
        ('POST /api/v1/backdrops/upload', 'Upload backdrop image file.'),
        ('GET /api/v1/projects/{project_id}/backdrops', 'Get all backdrops for a project.'),
        ('PUT /api/v1/backdrops/{backdrop_id}', 'Update a backdrop.'),
        ('DELETE /api/v1/backdrops/{backdrop_id}', 'Delete a backdrop.')
    ]
    
    for endpoint, description in backdrop_endpoints:
        para = doc.add_paragraph()
        run = para.add_run(f'{endpoint}: ')
        run.bold = True
        para.add_run(description)
    
    doc.add_heading('5.6 Stage Management Endpoints', 2)
    stage_endpoints = [
        ('GET /api/v1/projects/{project_id}/stage', 'Get stage settings for a project.'),
        ('PUT /api/v1/projects/{project_id}/stage', 'Update stage settings.'),
        ('PUT /api/v1/projects/{project_id}/stage/backdrop', 'Set active backdrop for stage.'),
        ('GET /api/v1/projects/{project_id}/stage/complete', 'Get complete stage data (settings, sprites, backdrops).')
    ]
    
    for endpoint, description in stage_endpoints:
        para = doc.add_paragraph()
        run = para.add_run(f'{endpoint}: ')
        run.bold = True
        para.add_run(description)
    
    doc.add_heading('5.7 Variable & List Endpoints', 2)
    variable_endpoints = [
        ('POST /api/v1/variables', 'Create a variable (global or sprite-specific).'),
        ('GET /api/v1/projects/{project_id}/variables', 'Get variables for a project or sprite.'),
        ('GET /api/v1/variables/{variable_id}', 'Get a specific variable.'),
        ('PUT /api/v1/variables/{variable_id}', 'Update a variable.'),
        ('DELETE /api/v1/variables/{variable_id}', 'Delete a variable.'),
        ('POST /api/v1/lists', 'Create a list (global or sprite-specific).'),
        ('GET /api/v1/projects/{project_id}/lists', 'Get lists for a project or sprite.'),
        ('PUT /api/v1/lists/{list_id}', 'Update a list.'),
        ('DELETE /api/v1/lists/{list_id}', 'Delete a list.')
    ]
    
    for endpoint, description in variable_endpoints:
        para = doc.add_paragraph()
        run = para.add_run(f'{endpoint}: ')
        run.bold = True
        para.add_run(description)
    
    doc.add_heading('5.8 Code Execution Endpoints', 2)
    execution_endpoints = [
        ('POST /api/v1/execute', 'Execute code in sandbox. Requires: code, language, timeout. Returns: ExecutionResult with output, error, execution_time.')
    ]
    
    for endpoint, description in execution_endpoints:
        para = doc.add_paragraph()
        run = para.add_run(f'{endpoint}: ')
        run.bold = True
        para.add_run(description)
    
    doc.add_heading('5.9 Library Endpoints', 2)
    library_endpoints = [
        ('GET /api/v1/library/sprites', 'Get available sprites from library. Supports category and search filters.'),
        ('GET /api/v1/library/backdrops', 'Get available backdrops from library.'),
        ('POST /api/v1/projects/{project_id}/library/sprite', 'Add a library sprite to project.'),
        ('POST /api/v1/projects/{project_id}/library/backdrop', 'Add a library backdrop to project.')
    ]
    
    for endpoint, description in library_endpoints:
        para = doc.add_paragraph()
        run = para.add_run(f'{endpoint}: ')
        run.bold = True
        para.add_run(description)
    
    doc.add_heading('5.10 Sharing & Collaboration Endpoints', 2)
    sharing_endpoints = [
        ('POST /api/v1/projects/{project_id}/share', 'Share project with other users. Defines permissions.')
    ]
    
    for endpoint, description in sharing_endpoints:
        para = doc.add_paragraph()
        run = para.add_run(f'{endpoint}: ')
        run.bold = True
        para.add_run(description)
    
    doc.add_heading('5.11 Health Check', 2)
    doc.add_paragraph('GET /health - Health check endpoint. Returns service status.')
    
    # 6. Authentication & Security
    doc.add_page_break()
    doc.add_heading('6. Authentication & Security', 1)
    
    doc.add_heading('6.1 Authentication Flow', 2)
    doc.add_paragraph(
        'The application uses JWT (JSON Web Tokens) for authentication. The flow is as follows:'
    )
    auth_flow = [
        'User registers with username, email, and password',
        'Password is hashed using bcrypt before storage',
        'User logs in with username and password',
        'System validates credentials and generates JWT token',
        'Token is returned to client (valid for 7 days by default)',
        'Client includes token in Authorization header: "Bearer <token>"',
        'Protected endpoints validate token and extract user information'
    ]
    
    for step in auth_flow:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('6.2 Security Features', 2)
    security_features = [
        'Password hashing using bcrypt with automatic salt generation',
        'JWT tokens with expiration (configurable, default 7 days)',
        'OAuth2 password flow for token generation',
        'Token validation on protected endpoints',
        'User role-based access control (Admin, Teacher, Student, User)',
        'Account activation status checking',
        'CORS middleware for cross-origin requests',
        'Input validation using Pydantic schemas',
        'SQL injection prevention via SQLAlchemy ORM',
        'Code execution sandboxing with timeout limits'
    ]
    
    for feature in security_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('6.3 Authorization', 2)
    doc.add_paragraph(
        'Authorization is enforced at multiple levels:'
    )
    authorization_levels = [
        'Project ownership: Only project owner can modify/delete projects',
        'Public projects: Can be viewed by anyone, but only owner can edit',
        'Shared projects: Access controlled by ProjectShare permissions',
        'Resource ownership: Sprites, costumes, backdrops inherit project permissions',
        'Role-based: Admin users have elevated privileges'
    ]
    
    for level in authorization_levels:
        doc.add_paragraph(level, style='List Bullet')
    
    # 7. Code Execution & Sandbox
    doc.add_page_break()
    doc.add_heading('7. Code Execution & Sandbox', 1)
    
    doc.add_heading('7.1 Sandbox Architecture', 2)
    doc.add_paragraph(
        'The code execution sandbox provides a safe environment for running user-generated code:'
    )
    sandbox_features = [
        'Isolated execution using Python subprocess',
        'Configurable timeout (default 10 seconds, max 60 seconds)',
        'Automatic cleanup of temporary files',
        'Capture of stdout and stderr',
        'Execution time measurement',
        'Exit code tracking',
        'Support for Python and JavaScript execution',
        'Optional Docker-based execution for better isolation (production)'
    ]
    
    for feature in sandbox_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('7.2 Code Safety', 2)
    doc.add_paragraph(
        'Basic code validation checks for dangerous patterns:'
    )
    dangerous_patterns = [
        'import os, sys, subprocess, shutil, socket',
        '__import__',
        'eval(), exec()',
        'open(), file()',
        'input(), raw_input()'
    ]
    
    doc.add_paragraph('Note: For production, Docker-based execution is recommended for better isolation.')
    
    doc.add_heading('7.3 Execution Logging', 2)
    doc.add_paragraph(
        'All code executions are logged with:'
    )
    log_fields = [
        'User ID and Project ID',
        'Code content and language',
        'Output and errors',
        'Execution time',
        'Exit code',
        'Timestamp'
    ]
    
    for field in log_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    # 8. Installation & Setup
    doc.add_page_break()
    doc.add_heading('8. Installation & Setup', 1)
    
    doc.add_heading('8.1 Prerequisites', 2)
    prerequisites = [
        'Python 3.8 or higher',
        'PostgreSQL database (or SQLite for development)',
        'pip package manager',
        'Virtual environment (recommended)',
        'Node.js (optional, for JavaScript execution)',
        'Docker (optional, for production code execution)'
    ]
    
    for prereq in prerequisites:
        doc.add_paragraph(prereq, style='List Bullet')
    
    doc.add_heading('8.2 Installation Steps', 2)
    installation_steps = [
        'Clone or download the project',
        'Create a virtual environment: python -m venv venv',
        'Activate virtual environment:',
        '  - Windows: venv\\Scripts\\activate',
        '  - Linux/Mac: source venv/bin/activate',
        'Install dependencies: pip install -r requirements.txt',
        'Configure database connection in database.py',
        'Set environment variables (if needed)',
        'Initialize database: python database.py',
        'Run the application: python main.py or uvicorn main:app --reload'
    ]
    
    for step in installation_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('8.3 Database Setup', 2)
    doc.add_paragraph(
        'The application supports multiple database backends:'
    )
    
    db_setup_code = """
# PostgreSQL (Production)
DATABASE_URL = "postgresql://user:password@host:port/dbname"

# MySQL
DATABASE_URL = "mysql+pymysql://user:password@host:port/dbname"

# SQLite (Development)
DATABASE_URL = "sqlite:///./blockly_platform.db"
"""
    
    add_code_block(doc, db_setup_code)
    
    doc.add_paragraph(
        'Database tables are automatically created on first run via SQLAlchemy metadata.'
    )
    
    # 9. Configuration
    doc.add_page_break()
    doc.add_heading('9. Configuration', 1)
    
    doc.add_heading('9.1 Environment Variables', 2)
    doc.add_paragraph(
        'Key configuration options:'
    )
    
    config_table = doc.add_table(rows=1, cols=2)
    config_table.style = 'Light Grid Accent 1'
    hdr_cells = config_table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Description'
    
    config_vars = [
        ('DATABASE_URL', 'Database connection string'),
        ('SECRET_KEY', 'JWT secret key (change in production!)'),
        ('ACCESS_TOKEN_EXPIRE_MINUTES', 'JWT token expiration time (default: 10080 = 7 days)'),
        ('CORS_ORIGINS', 'Allowed CORS origins (comma-separated)')
    ]
    
    for var, desc in config_vars:
        row_cells = config_table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = desc
    
    doc.add_heading('9.2 Security Configuration', 2)
    doc.add_paragraph(
        'Important security settings in auth.py:'
    )
    
    security_config = """
SECRET_KEY = "YOUR_SECRET_KEY_CHANGE_THIS_IN_PRODUCTION"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24 * 7  # 7 days
"""
    
    add_code_block(doc, security_config)
    
    doc.add_paragraph(
        '⚠️ WARNING: Change SECRET_KEY in production! Use a strong, random key.'
    )
    
    # 10. Usage Examples
    doc.add_page_break()
    doc.add_heading('10. Usage Examples', 1)
    
    doc.add_heading('10.1 User Registration', 2)
    registration_example = """
POST /api/v1/register
Content-Type: application/json

{
  "username": "johndoe",
  "email": "john@example.com",
  "password": "securepassword123",
  "full_name": "John Doe"
}
"""
    add_code_block(doc, registration_example)
    
    doc.add_heading('10.2 User Login', 2)
    login_example = """
POST /api/v1/token
Content-Type: application/x-www-form-urlencoded

username=johndoe&password=securepassword123

Response:
{
  "access_token": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",
  "token_type": "bearer"
}
"""
    add_code_block(doc, login_example)
    
    doc.add_heading('10.3 Create Project', 2)
    create_project_example = """
POST /api/v1/projects
Authorization: Bearer <token>
Content-Type: application/json

{
  "title": "My First Project",
  "description": "A simple visual programming project",
  "workspace_xml": "<xml>...</xml>",
  "code_language": "python",
  "is_public": false
}
"""
    add_code_block(doc, create_project_example)
    
    doc.add_heading('10.4 Create Sprite', 2)
    create_sprite_example = """
POST /api/v1/sprites
Authorization: Bearer <token>
Content-Type: application/json

{
  "project_id": 1,
  "name": "Cat",
  "x_position": 0.0,
  "y_position": 0.0,
  "direction": 90.0,
  "size": 100.0,
  "is_visible": true
}
"""
    add_code_block(doc, create_sprite_example)
    
    doc.add_heading('10.5 Execute Code', 2)
    execute_code_example = """
POST /api/v1/execute
Authorization: Bearer <token>
Content-Type: application/json

{
  "code": "print('Hello, World!')\\nfor i in range(5):\\n    print(i)",
  "language": "python",
  "timeout": 10,
  "project_id": 1
}

Response:
{
  "output": "Hello, World!\\n0\\n1\\n2\\n3\\n4",
  "error": null,
  "execution_time": 0.023,
  "exit_code": 0,
  "success": true
}
"""
    add_code_block(doc, execute_code_example)
    
    # 11. Development Guidelines
    doc.add_page_break()
    doc.add_heading('11. Development Guidelines', 1)
    
    doc.add_heading('11.1 Code Style', 2)
    doc.add_paragraph(
        'Follow PEP 8 Python style guide. The project uses:'
    )
    style_guidelines = [
        'Type hints for function parameters and return values',
        'Docstrings for all functions and classes',
        'Consistent naming: snake_case for variables/functions, PascalCase for classes',
        'Modular code organization: separate concerns (models, schemas, crud, auth)',
        'Error handling with appropriate HTTP status codes'
    ]
    
    for guideline in style_guidelines:
        doc.add_paragraph(guideline, style='List Bullet')
    
    doc.add_heading('11.2 API Design Principles', 2)
    api_principles = [
        'RESTful API design with clear resource naming',
        'Consistent endpoint patterns: /api/v1/{resource}',
        'Proper HTTP methods: GET (read), POST (create), PUT (update), DELETE (delete)',
        'Status codes: 200 (success), 201 (created), 400 (bad request), 401 (unauthorized), 403 (forbidden), 404 (not found), 500 (server error)',
        'Request/response validation using Pydantic schemas',
        'Comprehensive error messages',
        'Pagination for list endpoints (skip, limit)'
    ]
    
    for principle in api_principles:
        doc.add_paragraph(principle, style='List Bullet')
    
    doc.add_heading('11.3 Database Best Practices', 2)
    db_practices = [
        'Use SQLAlchemy ORM instead of raw SQL',
        'Define relationships properly with cascade options',
        'Use transactions for multi-step operations',
        'Index frequently queried fields',
        'Use soft deletes where appropriate (future enhancement)',
        'Version control for projects (already implemented)'
    ]
    
    for practice in db_practices:
        doc.add_paragraph(practice, style='List Bullet')
    
    # 12. Deployment
    doc.add_page_break()
    doc.add_heading('12. Deployment', 1)
    
    doc.add_heading('12.1 Production Considerations', 2)
    production_considerations = [
        'Use PostgreSQL for production database (not SQLite)',
        'Set strong SECRET_KEY as environment variable',
        'Configure proper CORS origins (not "*")',
        'Use HTTPS for all API communications',
        'Implement rate limiting',
        'Use Docker for code execution sandbox',
        'Set up proper logging and monitoring',
        'Configure database connection pooling',
        'Use environment variables for all sensitive configuration',
        'Set up backup strategy for database',
        'Implement proper error handling and logging',
        'Use reverse proxy (nginx) in front of application',
        'Configure firewall rules',
        'Set up SSL/TLS certificates'
    ]
    
    for consideration in production_considerations:
        doc.add_paragraph(consideration, style='List Bullet')
    
    doc.add_heading('12.2 Running with Uvicorn', 2)
    uvicorn_examples = """
# Development (with auto-reload)
uvicorn main:app --reload --host 0.0.0.0 --port 8000

# Production (with workers)
uvicorn main:app --host 0.0.0.0 --port 8000 --workers 4

# With SSL
uvicorn main:app --host 0.0.0.0 --port 8000 --ssl-keyfile key.pem --ssl-certfile cert.pem
"""
    add_code_block(doc, uvicorn_examples)
    
    doc.add_heading('12.3 Docker Deployment', 2)
    doc.add_paragraph(
        'Example Dockerfile:'
    )
    
    dockerfile_example = """
FROM python:3.11-slim

WORKDIR /app

COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

COPY . .

EXPOSE 8000

CMD ["uvicorn", "main:app", "--host", "0.0.0.0", "--port", "8000"]
"""
    add_code_block(doc, dockerfile_example)
    
    # Appendices
    doc.add_page_break()
    doc.add_heading('Appendices', 1)
    
    doc.add_heading('A. Dependencies', 2)
    doc.add_paragraph('Complete list from requirements.txt:')
    
    dependencies = [
        'fastapi==0.104.1',
        'uvicorn[standard]==0.24.0',
        'sqlalchemy==2.0.23',
        'python-jose[cryptography]==3.3.0',
        'passlib[bcrypt]==1.7.4',
        'pydantic==2.9.2',
        'email-validator==2.1.0',
        'python-multipart==0.0.6',
        'python-dotenv==1.0.0',
        'pytest==7.4.3',
        'black==23.11.0',
        'Pillow (for image processing)'
    ]
    
    for dep in dependencies:
        doc.add_paragraph(dep, style='List Bullet')
    
    doc.add_heading('B. API Response Formats', 2)
    doc.add_paragraph('Standard response formats:')
    
    response_formats = [
        ('Success Response', '200 OK with data object'),
        ('Created Response', '201 Created with created object'),
        ('Error Response', '400/401/403/404/500 with {"detail": "error message"}'),
        ('List Response', '200 OK with array of objects'),
        ('Empty Response', '204 No Content for DELETE operations')
    ]
    
    for name, desc in response_formats:
        para = doc.add_paragraph()
        run = para.add_run(f'{name}: ')
        run.bold = True
        para.add_run(desc)
    
    doc.add_heading('C. Common Error Codes', 2)
    error_codes = [
        ('400 Bad Request', 'Invalid input data or validation error'),
        ('401 Unauthorized', 'Missing or invalid authentication token'),
        ('403 Forbidden', 'User lacks permission for the operation'),
        ('404 Not Found', 'Resource not found'),
        ('500 Internal Server Error', 'Server-side error')
    ]
    
    for code, desc in error_codes:
        para = doc.add_paragraph()
        run = para.add_run(f'{code}: ')
        run.bold = True
        para.add_run(desc)
    
    # Footer
    doc.add_page_break()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('End of Documentation')
    run.italic = True
    run.font.size = Pt(12)
    
    # Save document
    output_file = 'PictoBlox_Python_Backend_Documentation.docx'
    doc.save(output_file)
    print(f'Documentation saved to {output_file}')
    return output_file

if __name__ == '__main__':
    create_documentation()

